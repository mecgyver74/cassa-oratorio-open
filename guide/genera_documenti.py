"""
Genera 4 documenti Cassa Dalila con la stessa formattazione del manuale originale.
Strategia: apre l'originale, conserva solo il frontespizio (paras 0-8),
aggiorna i testi del titolo, aggiunge il nuovo contenuto.
"""
import copy, sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt

ORIG      = r"C:\temp\manuale_orig.docx"
OUT       = r"G:\.shortcut-targets-by-id\1MeKnNdaDF-U2crMSnvSXGJEJrMvsOIGF\Cassa_Dalila"
GUIDE_OUT = OUT + r"\guide"

# ── helpers ────────────────────────────────────────────────────────────────────

def get_style_id(p):
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None: return 'Normal'
    ps = pPr.find(qn('w:pStyle'))
    return ps.get(qn('w:val')) if ps is not None else 'Normal'

def set_run_text(p, text):
    """Sostituisce tutti i run di un paragrafo col testo dato, preservando il formato del primo run."""
    fmt = {}
    if p.runs:
        r0 = p.runs[0]
        fmt = {'bold': r0.bold, 'italic': r0.italic, 'size': r0.font.size}
        try: fmt['color'] = r0.font.color.rgb
        except: fmt['color'] = None
    for r in p._element.findall(qn('w:r')):
        p._element.remove(r)
    run = p.add_run(text)
    if fmt.get('bold'):   run.bold   = True
    if fmt.get('italic'): run.italic = True
    if fmt.get('size'):   run.font.size  = fmt['size']
    if fmt.get('color'):  run.font.color.rgb = fmt['color']

def set_spacing(p, before=None, after=None):
    pPr = p._element.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing'); pPr.append(sp)
    if before is not None: sp.set(qn('w:before'), str(before))
    if after  is not None: sp.set(qn('w:after'),  str(after))

def add_h1(doc, tmpl_h1, text):
    new_el = copy.deepcopy(tmpl_h1._element)
    doc.element.body.append(new_el)
    for p in doc.paragraphs:
        if p._element is new_el:
            set_run_text(p, text)
            return p

def add_h2(doc, tmpl_h2, text):
    new_el = copy.deepcopy(tmpl_h2._element)
    doc.element.body.append(new_el)
    for p in doc.paragraphs:
        if p._element is new_el:
            set_run_text(p, text)
            return p

def add_normal(doc, tmpl_normal, text, after=120):
    new_el = copy.deepcopy(tmpl_normal._element)
    doc.element.body.append(new_el)
    for p in doc.paragraphs:
        if p._element is new_el:
            set_run_text(p, text)
            set_spacing(p, after=after)
            return p

def add_bullet(doc, tmpl_bullet, text):
    new_el = copy.deepcopy(tmpl_bullet._element)
    doc.element.body.append(new_el)
    for p in doc.paragraphs:
        if p._element is new_el:
            set_run_text(p, text)
            set_spacing(p, before=40, after=40)
            return p

def add_numbered(doc, tmpl_num, text):
    new_el = copy.deepcopy(tmpl_num._element)
    doc.element.body.append(new_el)
    for p in doc.paragraphs:
        if p._element is new_el:
            set_run_text(p, text)
            set_spacing(p, before=40, after=40)
            return p

QUOTE_COLORS = {
    'suggerimento': ('#16A34A', 'F0FDF4'),
    'attenzione':   ('#DC2626', 'FEF2F2'),
    'nota':         ('#D97706', 'FEF3C7'),
}

def add_quote(doc, tmpl_quote, text):
    """Clona un blockquote e aggiusta colore/sfondo in base alla parola chiave."""
    new_el = copy.deepcopy(tmpl_quote._element)
    doc.element.body.append(new_el)
    for p in doc.paragraphs:
        if p._element is new_el:
            # Aggiorna colore border e shading in base al tipo
            first_word = text.lstrip('*').split(':')[0].lower()
            color_hex, fill_hex = QUOTE_COLORS.get(first_word, ('#2563EB', 'EFF6FF'))
            pPr = p._element.get_or_add_pPr()
            # border
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                for side in pBdr:
                    side.set(qn('w:color'), color_hex.lstrip('#'))
            # shading
            shd = pPr.find(qn('w:shd'))
            if shd is not None:
                shd.set(qn('w:fill'), fill_hex)
            # clear runs and set bold label + rest
            for r in p._element.findall(qn('w:r')):
                p._element.remove(r)
            import re
            m = re.match(r'^([^:]+:)\s*(.*)', text)
            if m:
                label, rest = m.group(1), m.group(2)
                r = p.add_run(label + ' ')
                r.bold = True
                r.font.color.rgb = RGBColor(
                    int(color_hex[1:3], 16), int(color_hex[3:5], 16), int(color_hex[5:7], 16))
                p.add_run(rest)
            else:
                p.add_run(text)
            return p

def add_page_break(doc, tmpl_pgbrk):
    new_el = copy.deepcopy(tmpl_pgbrk._element)
    doc.element.body.append(new_el)

def add_table(doc, orig_table):
    """Clona una tabella dall'originale (struttura + bordi)."""
    new_tbl = copy.deepcopy(orig_table._element)
    doc.element.body.append(new_tbl)
    # trova la tabella clonata
    for t in doc.tables:
        if t._element is new_tbl:
            return t
    return None

def set_cell(cell, text, bold=False):
    for p in cell.paragraphs:
        for r in p._element.findall(qn('w:r')):
            p._element.remove(r)
    run = cell.paragraphs[0].add_run(text)
    if bold: run.bold = True

# ── funzione principale di generazione ────────────────────────────────────────

def make_doc(filename, big_title, subtitle1, main_title, subtitle2, version, content_fn, dest=None):
    """
    Crea un nuovo documento clonando il frontespizio dall'originale
    e aggiungendo il contenuto tramite content_fn(doc, templates).
    """
    doc = Document(ORIG)
    body = doc.element.body
    children = list(body)

    # Trova l'indice del section break (para[8]) nel body
    # I primi 9 paragrafi sono il frontespizio
    front_paras = doc.paragraphs[:9]

    # Elimina tutto ciò che viene dopo il frontespizio
    front_els = {p._element for p in front_paras}
    to_remove = []
    found_front_end = False
    for child in children:
        if found_front_end:
            to_remove.append(child)
        else:
            # Il section break è nel para[8]
            if child is front_paras[8]._element:
                found_front_end = True
    for el in to_remove:
        body.remove(el)

    # Aggiorna testi del frontespizio
    # para[1] = big title, para[2] = subtitle1, para[4] = main_title
    # para[5] = subtitle2, para[6] = version
    set_run_text(doc.paragraphs[1], big_title)
    set_run_text(doc.paragraphs[2], subtitle1)
    set_run_text(doc.paragraphs[4], main_title)
    set_run_text(doc.paragraphs[5], subtitle2)
    set_run_text(doc.paragraphs[6], version)

    # Template paragraphs dall'originale — usati come "stampi" per clonare
    # Dobbiamo prenderli dalla copia dell'originale PRIMA che venga eliminato
    # Usiamo il documento originale separato per i template
    orig = Document(ORIG)
    def find_orig(text_frag, style_id=None):
        for p in orig.paragraphs:
            if text_frag in p.text:
                if style_id is None: return p
                if get_style_id(p) == style_id: return p
        return None

    templates = {
        'h1':      find_orig('1. Introduzione', 'Heading10'),
        'h2':      find_orig('Installazione con installer', 'Heading20'),
        'bullet':  find_orig('Gestire il magazzino', 'Paragrafoelenco'),
        'num':     find_orig('Scarica e installa Node.js', 'Paragrafoelenco'),
        'normal':  find_orig('Il sistema e completamente autonomo'),
        'suggest': find_orig("Suggerimento: L'installer"),
        'warn':    find_orig('Attenzione: Non chiudere'),
        'note':    find_orig('Nota: Se il tablet'),
        'pgbrk':   orig.paragraphs[7],   # page break
        'table_light': orig.tables[0],    # tabella con header grigio chiaro
        'table_dark':  orig.tables[10],   # tabella con header navy
    }

    # Funzioni di aggiunta legate a doc e templates
    def h1(text):    return add_h1(doc, templates['h1'], text)
    def h2(text):    return add_h2(doc, templates['h2'], text)
    def norm(text, after=120): return add_normal(doc, templates['normal'], text, after)
    def blt(text):   return add_bullet(doc, templates['bullet'], text)
    def num(text):   return add_numbered(doc, templates['num'], text)
    def sug(text):   return add_quote(doc, templates['suggest'], text)
    def warn(text):  return add_quote(doc, templates['warn'], text)
    def note(text):  return add_quote(doc, templates['note'], text)
    def pgbrk():     return add_page_break(doc, templates['pgbrk'])

    def tbl_light(header_row, data_rows):
        t = add_table(doc, templates['table_light'])
        if t is None: return
        # Header
        while len(t.columns) < len(header_row):
            pass  # usa la struttura clonata
        # Resize: remove extra rows, keep 1 header
        while len(t.rows) > 1:
            t._tbl.remove(t.rows[-1]._tr)
        for ci, h in enumerate(header_row[:len(t.columns)]):
            set_cell(t.cell(0, ci), h, bold=True)
        for row_data in data_rows:
            row = t.add_row()
            for ci, v in enumerate(row_data[:len(t.columns)]):
                set_cell(row.cells[ci], v)

    def tbl_dark(header_row, data_rows):
        t = add_table(doc, templates['table_dark'])
        if t is None: return
        while len(t.rows) > 1:
            t._tbl.remove(t.rows[-1]._tr)
        for ci, h in enumerate(header_row[:len(t.columns)]):
            set_cell(t.cell(0, ci), h, bold=True)
        for row_data in data_rows:
            row = t.add_row()
            for ci, v in enumerate(row_data[:len(t.columns)]):
                set_cell(row.cells[ci], v)

    # Genera il contenuto
    content_fn(h1, h2, norm, blt, num, sug, warn, note, pgbrk, tbl_light, tbl_dark)

    path = f"{dest or OUT}\\{filename}"
    doc.save(path)
    print(f"✓ {filename}")
    return path

# ── documento compatto (senza frontespizio) ───────────────────────────────────

def make_compact_doc(filename, title, subtitle, content_fn, dest=None):
    """
    Crea un documento compatto senza frontespizio: solo un'intestazione
    breve e il contenuto. Ideale per guide rapide su 1-2 pagine.
    """
    doc = Document(ORIG)
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl', 'sdt'):
            body.remove(child)

    orig = Document(ORIG)
    def find_orig(text_frag, style_id=None):
        for p in orig.paragraphs:
            if text_frag in p.text:
                if style_id is None: return p
                if get_style_id(p) == style_id: return p
        return None

    tmpl = {
        'h1':      find_orig('1. Introduzione', 'Heading10'),
        'h2':      find_orig('Installazione con installer', 'Heading20'),
        'bullet':  find_orig('Gestire il magazzino', 'Paragrafoelenco'),
        'num':     find_orig('Scarica e installa Node.js', 'Paragrafoelenco'),
        'normal':  find_orig('Il sistema e completamente autonomo'),
        'suggest': find_orig("Suggerimento: L'installer"),
        'warn':    find_orig('Attenzione: Non chiudere'),
        'note':    find_orig('Nota: Se il tablet'),
        'table_light': orig.tables[0],
    }

    # Intestazione compatta
    tp = doc.add_paragraph()
    r1 = tp.add_run(title)
    r1.bold = True; r1.font.size = Pt(18)
    r1.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
    set_spacing(tp, before=0, after=60)

    if subtitle:
        sp2 = doc.add_paragraph()
        r2 = sp2.add_run(subtitle)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        set_spacing(sp2, before=0, after=200)

    # Helper con spaziatura ridotta
    def h1(text):
        p = add_h1(doc, tmpl['h1'], text)
        set_spacing(p, before=160, after=60)
        return p

    def h2(text):
        p = add_h2(doc, tmpl['h2'], text)
        set_spacing(p, before=100, after=40)
        return p

    def norm(text, after=80):
        return add_normal(doc, tmpl['normal'], text, after)

    def blt(text):
        p = add_bullet(doc, tmpl['bullet'], text)
        set_spacing(p, before=20, after=20)
        return p

    def num(text):
        p = add_numbered(doc, tmpl['num'], text)
        set_spacing(p, before=20, after=20)
        return p

    def sug(text):  return add_quote(doc, tmpl['suggest'], text)
    def warn(text): return add_quote(doc, tmpl['warn'], text)
    def note(text): return add_quote(doc, tmpl['note'], text)

    def tbl_light(header_row, data_rows):
        t = add_table(doc, tmpl['table_light'])
        if t is None: return
        while len(t.rows) > 1:
            t._tbl.remove(t.rows[-1]._tr)
        for ci, h_text in enumerate(header_row[:len(t.columns)]):
            set_cell(t.cell(0, ci), h_text, bold=True)
        for row_data in data_rows:
            row = t.add_row()
            for ci, v in enumerate(row_data[:len(t.columns)]):
                set_cell(row.cells[ci], v)

    content_fn(h1, h2, norm, blt, num, sug, warn, note, tbl_light)

    path = f"{dest or OUT}\\{filename}"
    doc.save(path)
    print(f"✓ {filename}")
    return path


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 1 — CHANGELOG
# ══════════════════════════════════════════════════════════════════════════════

def content_changelog(h1, h2, norm, blt, num, sug, warn, note, pgbrk, tbl_light, tbl_dark):

    h1('1. Nuovi metodi di pagamento')
    h2('Satispay')
    norm('È stato aggiunto Satispay come metodo di pagamento, accanto a Contanti, Carta e Omaggio.')
    blt('Selezionabile nel pannello di pagamento come gli altri metodi')
    blt('Tracciato separatamente nelle statistiche e nel report di chiusura')
    blt('Incluso nell\'incasso reale (soldi fisicamente ricevuti)')

    h2('Buono Volontario')
    norm('I volontari possono ora pagare con un buono consumazione precaricato. Vedi sezione dedicata.')

    pgbrk()
    h1('2. Metodo di pagamento obbligatorio')
    norm('In precedenza era possibile completare uno scontrino senza selezionare un metodo di pagamento.')
    norm('Ora il metodo è obbligatorio: il pulsante Conferma rimane disattivato finché non se ne sceglie uno. Questo garantisce che il riepilogo di fine serata sia sempre accurato.')

    pgbrk()
    h1('3. Gestione Buoni Volontari')
    h2('Configurazione')
    norm('In Setup → Volontari è possibile creare l\'elenco dei volontari con il valore del loro buono consumazione (es. € 5,00).')
    num('Vai su Setup → Volontari')
    num('Clicca + Nuovo')
    num('Inserisci nome, cognome, valore del buono ed eventuali note')
    num('Salva')

    h2('Utilizzo alla cassa')
    norm('Nel piede dello scontrino compare la sezione Buono Volontario:')
    num('Seleziona il volontario dal menu a tendina')
    num('Il sistema mostra il saldo residuo del buono per la sessione corrente')
    num('Inserisci l\'importo da scalare dal buono')
    num('Il totale da pagare si aggiorna in tempo reale')
    sug('Suggerimento: Se il buono copre l\'intero importo, il pagamento viene completato automaticamente. Se copre solo una parte, il residuo si paga con contanti, carta o Satispay.')

    h2('Azzeramento automatico')
    norm('Il saldo del buono si azzera a ogni chiusura cassa. All\'inizio di ogni nuova sessione i volontari ripartono con il buono pieno, indipendentemente dall\'uso precedente.')
    note('Nota: I buoni non si accumulano tra sessioni diverse e non sono trasferibili.')

    pgbrk()
    h1('4. Correzione contabile — Incasso reale vs Valore venduto')
    norm('I buoni non sono soldi che entrano in cassa: rappresentano un costo per l\'oratorio. La distinzione è ora esplicita in tutti i pannelli.')
    tbl_light(
        ['Voce', 'Cosa comprende'],
        [
            ['Incasso reale',        'Contanti + Carta + Satispay (soldi fisicamente incassati)'],
            ['Buoni (costo oratorio)','Valore dei buoni usati — mostrato separatamente, non è un\'entrata'],
            ['Valore venduto totale', 'Incasso reale + Buoni + Omaggi (tutto ciò che è uscito dal banco)'],
        ]
    )
    norm('Questa distinzione appare in: anteprima chiusura cassa, Statistiche, Dashboard, report XLS.')

    pgbrk()
    h1('5. Piccoli miglioramenti')
    h2('Numero tavolo nello scontrino')
    norm('È possibile assegnare un numero tavolo direttamente nello scontrino, cliccando il pulsante Tavolo nel piede.')

    h2('Modifica tipo di pagamento dallo Storico')
    norm('Dallo storico scontrini è ora possibile correggere il metodo di pagamento di uno scontrino già emesso, senza stornarlo e rifarlo.')

    h2('Filtro per metodo di pagamento nelle Statistiche')
    norm('In cima alla pagina Statistiche sono disponibili pulsanti rapidi per filtrare gli scontrini per tipo di pagamento: Tutti · Contanti · Carta · Satispay · Buono · Omaggio.')

    h2('Ordinamento volontari alfabetico')
    norm('Il menu di selezione del volontario alla cassa mostra i nomi in ordine alfabetico.')

    pgbrk()
    h1('6. Report di chiusura migliorato')
    norm('Il report di chiusura cassa è stato aggiornato. L\'email ora include un allegato Excel (.xlsx) al posto del vecchio file CSV.')
    tbl_light(
        ['Foglio XLS', 'Contenuto'],
        [
            ['Riepilogo',  'Incasso reale, buoni (costo), scontrini, media, contanti, carta, Satispay'],
            ['Venduto',    'Quantità e valore per ogni prodotto, inclusi gli omaggi'],
            ['Scontrini',  'Elenco completo con numero, data, lordo, buono, netto, pagamento, stato'],
        ]
    )
    sug('Suggerimento: Il file XLS si apre direttamente con Excel o LibreOffice Calc senza conversioni.')

    pgbrk()
    h1('7. Riepilogo tecnico')
    norm('Per chi gestisce l\'installazione:')
    blt('Nuova collezione volontari nel database')
    blt('Nuova collezione sessioni_cassa per lo storico delle chiusure')
    blt('Nuovo campo tipo_pagamento: aggiunto il valore "satispay" e "buono"')
    blt('Hook server-side aggiornato: genera XLS invece di CSV, calcola venduto per prodotto')
    blt('Migration automatica: eseguita al primo avvio del server aggiornato')
    warn('Attenzione: Prima di aggiornare il server esegui sempre un backup della cartella app\\pb_data\\.')


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 2 — GUIDA RAPIDA CASSA (compatta, 2 facciate, senza frontespizio)
# ══════════════════════════════════════════════════════════════════════════════

def content_guida_cassa(h1, h2, norm, blt, num, sug, warn, note, tbl_light):

    h1('1. Scontrino base — Contanti · Carta · Satispay')
    num('Aggiungi prodotti; usa + e − per le quantità')
    num('Ingredienti personalizzati (se disponibili): clicca 🔧 sulla riga → deseleziona → Conferma')
    num('Tavolo (opz.): piede scontrino → Tavolo → inserisci numero')
    num('Asporto (opz.): attiva il toggle Asporto nel piede (diventa arancione)')
    num('Clicca PAGA → scegli il metodo: Contanti · Carta · Satispay')
    num('Contanti: inserisci importo ricevuto — il resto si calcola; lascia vuoto per importo esatto')
    num('Clicca Conferma')

    h1('2. Omaggio')
    h2('Tutto lo scontrino')
    num('Componi scontrino → PAGA → Omaggio → Conferma (totale = €0)')
    h2('Singola riga')
    num('Clicca l\'icona omaggio direttamente sulla riga del prodotto nello scontrino')

    h1('3. Buono Volontario')
    num('Componi lo scontrino')
    num('Sezione Buono Volontario: seleziona il volontario — il sistema scala automaticamente fino al valore totale dello scontrino')
    num('Clicca PAGA')
    num('Se rimane un residuo: seleziona il metodo (Contanti · Carta · Satispay) → Conferma')
    note('Nota: Il saldo si azzera a ogni chiusura cassa — il volontario riparte sempre con il buono pieno.')

    h1('4. Correggere uno scontrino già emesso')
    blt('Rimuovere riga: Storico → scontrino → pulsante elimina sulla riga')
    blt('Modificare quantità: Storico → scontrino → +/− sulla riga')
    blt('Correggere metodo pagamento: Storico → scontrino → icona modifica sul tipo di pagamento')
    warn('Attenzione: Le modifiche sono registrate e non eliminano i dati originali.')

    h1('5. Stornare uno scontrino')
    num('Storico → trova lo scontrino → pulsante Storna → Conferma')
    norm('Lo scontrino viene marcato come stornato e non conta nelle statistiche.')
    sug('Suggerimento: Preferisci correggere la singola riga dallo Storico invece di stornare tutto.')


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 3 — GUIDA RAPIDA CHIUSURA CASSA (compatta, 2 facciate)
# ══════════════════════════════════════════════════════════════════════════════

def content_guida_chiusura(h1, h2, norm, blt, num, sug, warn, note, tbl_light):

    h1('1. Procedura di chiusura')
    norm('Solo gli amministratori possono eseguirla. È irreversibile — i dati non vengono cancellati.')
    num('Vai in Statistiche → pulsante 🔒 Chiudi cassa')
    num('Si apre l\'anteprima con il riepilogo della sessione')
    num('Verifica i totali (vedi sezione 2)')
    num('Assegna un nome alla sessione (opzionale — di default usa la data)')
    num('Clicca Chiudi cassa → Conferma')
    num('Attendi il completamento; al termine appare il riepilogo finale')
    warn('Attenzione: Prima di chiudere assicurati che nessun altro dispositivo abbia scontrini in corso.')

    h1('2. Cosa verificare nel riepilogo')
    tbl_light(
        ['Voce', 'Cosa fare'],
        [
            ['Incasso reale (Contanti + Carta + Satispay)', 'Confronta contanti col cassetto, carta col giornaliero POS'],
            ['Buoni (costo oratorio)',                      'Valore consumato dai volontari — non è denaro in cassa'],
            ['Omaggi',                                     'Valore articoli dati gratis'],
            ['Scontrini stornati',                         'Verificare se inattesi (mostrati in rosso)'],
        ]
    )
    sug('Suggerimento: Se i totali non tornano, controlla lo Storico prima di confermare.')

    h1('3. Report XLS via email')
    norm('Se configurata in Setup → Notifiche, l\'email arriva automaticamente con allegato Excel (3 fogli):')
    blt('Riepilogo — incasso reale, buoni, scontrini, media, per metodo di pagamento')
    blt('Venduto — quantità e valore per ogni prodotto, omaggi separati')
    blt('Scontrini — elenco completo della sessione')
    norm('I file vengono salvati anche localmente nella cartella chiusure\\.')

    h1('4. Dopo la chiusura')
    blt('Contatore scontrini azzera e riparte da #0001')
    blt('Saldi buoni volontari azzerati — ogni volontario riparte con il buono pieno')
    blt('Statistiche "Corrente" mostrano zero')
    blt('Sessione chiusa consultabile in Statistiche → Sessioni archiviate → seleziona dal menu')

    h1('5. Backup')
    num('Chiudi il browser')
    num('Doppio clic su BACKUP_DATI.bat nella cartella del programma')
    num('Il file viene salvato in backup\\ con data e ora nel nome')
    warn('Attenzione: Non cancellare mai la cartella app\\pb_data\\ — contiene tutto il database.')


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 4 — GUIDA RAPIDA MAGAZZINO (compatta, 2 facciate)
# ══════════════════════════════════════════════════════════════════════════════

def content_guida_magazzino(h1, h2, norm, blt, num, sug, warn, note, tbl_light):

    h1('1. Come funziona il magazzino')
    norm('Le scorte si aggiornano automaticamente ad ogni scontrino. La pagina Magazzino ha tre schede: Magazzini comuni, Prodotti singoli, Movimenti.')
    tbl_light(
        ['Evento', 'Effetto sulle scorte'],
        [
            ['Scontrino emesso',          'Scarico automatico dei prodotti venduti'],
            ['Scontrino stornato',        'Carico automatico (scorte ripristinate)'],
            ['Riga rimossa dallo storico','Carico automatico della quantità rimossa'],
            ['Carico manuale (+)',        'Aggiunge alla scorta attuale'],
            ['Rettifica (=)',             'Imposta la scorta al valore esatto'],
        ]
    )
    note('Nota: Scorta -1 = illimitata, non viene scalata dalle vendite.')

    h1('2. Carico e rettifica')
    h2('Carico + (merce ricevuta)')
    num('Magazzino → trova prodotto o magazzino comune → Carico (+)')
    num('Inserisci la quantità ricevuta → aggiungi nota opzionale → Salva')
    norm('Risultato: scorta precedente + quantità inserita.')
    sug('Suggerimento: Fai i carichi a inizio serata, prima di aprire la cassa.')
    h2('Rettifica = (conteggio fisico)')
    num('Magazzino → trova prodotto → Rettifica (=)')
    num('Inserisci il valore contato fisicamente → aggiungi nota → Salva')
    norm('Risultato: la scorta viene sovrascritta con il valore inserito.')
    warn('Attenzione: La rettifica sovrascrive il valore attuale — usala solo dopo un conteggio fisico.')

    h1('3. Magazzini comuni')
    norm('Permettono a più prodotti di condividere lo stesso stock (es. "Panini" usato da Hamburger, Hamburger bacon, Veggie burger).')
    h2('Creare un magazzino comune')
    num('Setup → Magazzini comuni → Nuovo → inserisci nome e scorta iniziale → Salva')
    h2('Collegare un prodotto')
    num('Setup → Prodotti → seleziona prodotto → campo Magazzino comune → Salva')
    norm('Ogni vendita del prodotto scarica il magazzino comune, non la scorta singola.')

    h1('4. Allarmi scorte basse')
    norm('Configura la soglia in Setup → Prodotti → campo Soglia allarme. Quando la scorta scende sotto soglia, il prodotto appare evidenziato nella pagina Magazzino e il bottone in cassa lampeggia con un alone rosso.')


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO 5 — GUIDA RAPIDA COMANDE (compatta, 2 facciate)
# ══════════════════════════════════════════════════════════════════════════════

def content_guida_comande(h1, h2, norm, blt, num, sug, warn, note, tbl_light):

    h1('1. Come funziona la schermata Comande')
    norm('Ogni scontrino confermato alla cassa genera automaticamente una comanda. Il personale di cucina o banco gestisce le comande da questa pagina, senza accedere alla cassa.')
    blt('Le comande appaiono in ordine di arrivo, dalla più vecchia alla più recente')
    blt('Ogni card mostra: numero scontrino, tavolo (se impostato), prodotti, ora di arrivo')
    blt('Gli ordini da asporto sono evidenziati in arancione')
    note('Nota: La pagina si aggiorna automaticamente — non serve ricaricare il browser.')

    h1('2. Gestire una comanda')
    num('Trova la comanda nella lista')
    num('Man mano che i prodotti sono pronti, selezionali sulla card (spunta singola riga)')
    num('Quando tutto è pronto, premi il pulsante Segna come evasa')
    num('La comanda scompare dalla lista attiva e viene archiviata')
    sug('Suggerimento: Puoi spuntare i prodotti uno alla volta mentre li prepari, senza evasione immediata.')

    h1('3. Ordini da asporto')
    blt('Evidenziati in arancione con dicitura ASPORTO visibile sulla card')
    blt('Seguono lo stesso flusso delle comande normali: spunta prodotti → Segna come evasa')

    h1('4. Filtri, riepilogo e ordinamento')
    h2('Filtro per famiglia')
    norm('In cima alla schermata sono presenti i pulsanti delle famiglie di prodotti (es. Cucina, Bar, Dolci). Selezionando una famiglia si vedono solo le comande che contengono prodotti di quella famiglia.')
    h2('Riepilogo')
    norm('Il pulsante Riepilogo mostra un conteggio aggregato di tutti i prodotti ancora da preparare, raggruppati per tipo. Utile per avere un colpo d\'occhio su quante porzioni mancano prima di servire.')
    h2('Ordinamento')
    norm('È possibile cambiare l\'ordinamento delle card tra: per ora di arrivo (default) e per tavolo.')


# ══════════════════════════════════════════════════════════════════════════════
# GENERA TUTTI I DOCUMENTI
# ══════════════════════════════════════════════════════════════════════════════

make_doc(
    'Changelog_Maggio_2026.docx',
    'CASSA DALILA',
    'Changelog e note di rilascio',
    'NOVITÀ DA MAGGIO 2026',
    'Aggiornamenti dalla versione 1.3 alla 1.4',
    'Maggio 2026',
    content_changelog
)

make_compact_doc(
    'Guida_Rapida_Cassa.docx',
    'GUIDA RAPIDA — OPERAZIONI DI CASSA',
    'Cassa Dalila v1.4  ·  Oratorio Dalila',
    content_guida_cassa,
    dest=GUIDE_OUT
)

make_compact_doc(
    'Guida_Rapida_Chiusura_Cassa.docx',
    'GUIDA RAPIDA — CHIUSURA CASSA',
    'Cassa Dalila v1.4  ·  Oratorio Dalila',
    content_guida_chiusura,
    dest=GUIDE_OUT
)

make_compact_doc(
    'Guida_Rapida_Magazzino.docx',
    'GUIDA RAPIDA — GESTIONE MAGAZZINO',
    'Cassa Dalila v1.4  ·  Oratorio Dalila',
    content_guida_magazzino,
    dest=GUIDE_OUT
)

make_compact_doc(
    'Guida_Rapida_Comande.docx',
    'GUIDA RAPIDA — GESTIONE COMANDE',
    'Cassa Dalila v1.4  ·  Oratorio Dalila',
    content_guida_comande,
    dest=GUIDE_OUT
)

print('\nTutti i documenti generati.')
