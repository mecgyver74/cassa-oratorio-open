"""
Aggiorna Manuale_Cassa_Dalila.docx v1.3 → v1.4 con chirurgia XML mirata.
Parte dall'originale: pagine, bordi, frontespizio rimangono intatti.
"""
import copy, re, sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor

SRC = r"C:\temp\manuale_orig.docx"
DST = r"C:\temp\Manuale_Cassa_Dalila_v14.docx"

doc = Document(SRC)

# ── helpers ───────────────────────────────────────────────────────────────────

def find(text_fragment, style_id=None):
    """Trova il primo paragrafo che contiene text_fragment."""
    for p in doc.paragraphs:
        if text_fragment in p.text:
            if style_id is None:
                return p
            pPr = p._element.find(qn('w:pPr'))
            ps = pPr.find(qn('w:pStyle')) if pPr is not None else None
            sid = ps.get(qn('w:val')) if ps is not None else 'Normal'
            if sid == style_id:
                return p
    return None

def set_text(p, new_text):
    """Sostituisce il testo di un paragrafo preservando lo stile del primo run."""
    # salva formato del primo run
    fmt = {}
    if p.runs:
        r0 = p.runs[0]
        fmt['bold']   = r0.bold
        fmt['italic'] = r0.italic
        fmt['size']   = r0.font.size
        try: fmt['color'] = r0.font.color.rgb
        except: fmt['color'] = None
    # rimuovi tutti i run
    for r in p._element.findall(qn('w:r')):
        p._element.remove(r)
    # aggiungi nuovo run
    r = p.add_run(new_text)
    if fmt.get('bold'):   r.bold   = True
    if fmt.get('italic'): r.italic = True
    if fmt.get('size'):   r.font.size = fmt['size']
    if fmt.get('color'):  r.font.color.rgb = fmt['color']
    return p

def insert_after(ref_p, style_template_p, text=''):
    """Inserisce paragrafo dopo ref_p, copiando stile da style_template_p."""
    from docx.text.paragraph import Paragraph
    new_el = copy.deepcopy(style_template_p._element)
    ref_p._element.addnext(new_el)
    p = Paragraph(new_el, ref_p._p.getparent() if hasattr(ref_p._p, 'getparent') else None)
    # trova nella lista live per avere il binding corretto
    for lp in doc.paragraphs:
        if lp._element is new_el:
            if text: set_text(lp, text)
            return lp
    # fallback
    if text: set_text(p, text)
    return p

def delete_para(p):
    p._element.getparent().remove(p._element)

def get_table_after(p):
    """Ritorna la prima tabella dopo il paragrafo p nel body XML."""
    body = doc.element.body
    children = list(body)
    idx = children.index(p._element)
    for child in children[idx+1:]:
        if child.tag.endswith('}tbl'):
            for t in doc.tables:
                if t._element is child:
                    return t
    return None

def set_cell_text(cell, text):
    for p in cell.paragraphs:
        for r in p._element.findall(qn('w:r')):
            p._element.remove(r)
    cell.paragraphs[0].add_run(text)

# ── template paragraphs (per clonare lo stile) ───────────────────────────────
# Questi esistono nell'originale e sono usati come "stampi" per nuovi paragrafi

tmpl_h2       = find('Installazione con installer', 'Heading20')
tmpl_bullet   = find('Gestire il magazzino', 'Paragrafoelenco')
tmpl_numbered = find('Scarica e installa Node.js da nodejs.org (scegli la versione LTS).', 'Paragrafoelenco')
tmpl_normal   = find('Il sistema e completamente autonomo')
tmpl_suggest  = find("Suggerimento: L'installer")
tmpl_warn     = find('Attenzione: Non chiudere')
tmpl_note     = find('Nota: Se il tablet')

# ── 1. Versione ───────────────────────────────────────────────────────────────
p = find('Versione 1.3')
if p: set_text(p, 'Versione 1.4  |  2026')
print("✓ Versione aggiornata")

# ── 2. Cosa puoi fare — aggiorna e aggiunge voci ─────────────────────────────
p15 = find('contanti, carta o omaggio', 'Paragrafoelenco')
if p15: set_text(p15, 'Vendere prodotti e incassare pagamenti in contanti, carta, Satispay, buono volontario o omaggio')

# Inserisci "Gestire buoni" dopo p15
p_buoni_feat = insert_after(p15, tmpl_bullet,
    'Gestire buoni per i volontari con saldo residuo in tempo reale')

p_stat = find('Consultare statistiche di vendita e incassi', 'Paragrafoelenco')
if p_stat: set_text(p_stat, 'Consultare statistiche di vendita e incassi con distinzione tra incasso reale e costi')

p_report = find('Stampare o esportare i dati in Excel', 'Paragrafoelenco')
if p_report: set_text(p_report, 'Stampare o esportare i dati in Excel')

# Aggiunge "Ricevere report via email" dopo "Stampare o esportare"
p_email_feat = insert_after(p_report, tmpl_bullet,
    'Ricevere il report di chiusura via email con allegato Excel multi-foglio')

print("✓ 'Cosa puoi fare' aggiornata")

# ── 3. Setup — aggiunge sezione Volontari dopo Comande ───────────────────────
p_attenzione_comande = find('Attenzione: Se non assegni nessuna famiglia', 'Normal')
p_aspetto_h2 = find('Aspetto dei pulsanti', 'Heading20')

# Inserisci PRIMA di "Aspetto dei pulsanti" (che segue già Comande):
# lavoriamo al contrario: cloniamo e inseriamo DOPO p_attenzione_comande

# H2 "Volontari"
p_vol_h2 = insert_after(p_attenzione_comande, tmpl_h2, 'Volontari (nuovo in v1.4)')
# Testo intro
p_vol_intro = insert_after(p_vol_h2, tmpl_normal,
    'Gestisci l\'elenco dei volontari e il valore del loro buono consumazione.')
# Steps
p_v1 = insert_after(p_vol_intro, tmpl_numbered, 'Vai su Setup → Volontari')
p_v2 = insert_after(p_v1,        tmpl_numbered, 'Clicca su + Nuovo')
p_v3 = insert_after(p_v2,        tmpl_numbered, 'Inserisci: nome e cognome, valore del buono (€), note opzionali')
p_v4 = insert_after(p_v3,        tmpl_numbered, 'Salva')
# Nota
p_vol_note = insert_after(p_v4, tmpl_note,
    'Nota: Il valore del buono si azzera automaticamente a ogni chiusura cassa — ogni serata i volontari ripartono con il buono pieno. I buoni non si accumulano tra sessioni.')

print("✓ Sezione 'Volontari' inserita")

# ── 4. "Come fare uno scontrino" — aggiorna steps ────────────────────────────
p_s1 = find('Seleziona il filtro di famiglia in alto', 'Paragrafoelenco')
if p_s1: set_text(p_s1, 'Seleziona il filtro di famiglia in alto per filtrare i prodotti')

p_s3 = find('Modifica le quantita con i pulsanti + e -', 'Paragrafoelenco')
if p_s3: set_text(p_s3, 'Modifica le quantità con i pulsanti + e − nello scontrino')

p_s4 = find('Per aggiungere una nota libera', 'Paragrafoelenco')
if p_s4: set_text(p_s4, 'Per aggiungere una nota a una riga, fai doppio clic sul nome del prodotto')

# Rimuovi step sconto (era punto 5 originale: "Applica uno sconto se necessario")
p_sconto_step = find('Applica uno sconto se necessario', 'Paragrafoelenco')
if p_sconto_step: delete_para(p_sconto_step)

# Para "Per assegnare un numero tavolo" diventa list item
p_tavolo = find('Per assegnare un numero tavolo')
if p_tavolo:
    # Clona struttura da un bullet numerato e aggiorna
    new_el = copy.deepcopy(tmpl_numbered._element)
    p_tavolo._element.getparent().replace(p_tavolo._element, new_el)
    for p in doc.paragraphs:
        if p._element is new_el:
            set_text(p, 'Per assegnare un numero tavolo, clicca il pulsante Tavolo nello scontrino')
            p_tavolo = p
            break

# Para "Per un ordine da asporto" diventa list item
p_asporto = find('Per un ordine da asporto, attiva il toggle Asporto')
if p_asporto:
    new_el = copy.deepcopy(tmpl_numbered._element)
    p_asporto._element.getparent().replace(p_asporto._element, new_el)
    for p in doc.paragraphs:
        if p._element is new_el:
            set_text(p, 'Per un ordine da asporto, attiva il toggle Asporto')
            p_asporto = p
            break

# Aggiunge step "Se c'è un volontario, selezionalo" prima di PAGA
p_paga = find('Clicca PAGA per aprire il pannello', 'Paragrafoelenco')
if p_paga:
    # Inserisci PRIMA di PAGA: aggiungi dopo p_asporto
    p_buono_step = insert_after(p_asporto, tmpl_numbered,
        'Se è presente un volontario, selezionalo nella sezione Buono Volontario (vedi sotto)')
    set_text(p_paga, 'Clicca PAGA per aprire il pannello di pagamento')

print("✓ 'Come fare uno scontrino' aggiornato")

# ── 5. Inserisci sezione "Buono Volontario" dopo "Come fare uno scontrino" ──
p_suggest_cassa = find('Usa il pulsante', 'Normal')   # il Suggerimento dopo scontrino
p_pagamento_h2  = find('Pagamento', 'Heading20')

# Inserisci dopo il Suggerimento (prima di "Pagamento")
anchor = p_suggest_cassa if p_suggest_cassa else p_paga

p_bv_h2 = insert_after(anchor, tmpl_h2, 'Buono Volontario (nuovo in v1.4)')
p_bv_1  = insert_after(p_bv_h2, tmpl_normal,
    'I volontari hanno diritto a un buono consumazione per la serata. Per utilizzarlo:')
p_bv_s1 = insert_after(p_bv_1, tmpl_numbered,
    'Nel piede dello scontrino, trova la sezione Buono Volontario')
p_bv_s2 = insert_after(p_bv_s1, tmpl_numbered,
    'Seleziona il volontario dal menu a tendina')
p_bv_s3 = insert_after(p_bv_s2, tmpl_numbered,
    'Viene mostrato automaticamente il saldo residuo del buono per questa sessione')
p_bv_s4 = insert_after(p_bv_s3, tmpl_numbered,
    'Inserisci l\'importo da scalare dal buono (non può superare il saldo disponibile né il totale)')
p_bv_s5 = insert_after(p_bv_s4, tmpl_numbered,
    'Il totale da pagare si aggiorna in tempo reale')
p_bv_note = insert_after(p_bv_s5, tmpl_suggest,
    'Suggerimento: Se il buono copre l\'intero importo, il pagamento viene completato senza richiedere un metodo aggiuntivo. Se copre solo una parte, il residuo si paga con contanti, carta o Satispay.')

print("✓ Sezione 'Buono Volontario' inserita")

# ── 6. Aggiorna sezione "Pagamento" ──────────────────────────────────────────
p_pag_intro = find('Nel pannello di pagamento puoi scegliere il metodo')
if p_pag_intro:
    set_text(p_pag_intro,
        'Cliccando PAGA si apre il pannello di pagamento. '
        'Il metodo di pagamento deve essere selezionato esplicitamente — '
        'il pulsante Conferma rimane disattivato finché non si sceglie.')

# Aggiorna tabella pagamento (la tabella immediatamente dopo il paragrafo intro)
t_pag = get_table_after(p_pag_intro)
if t_pag and len(t_pag.rows) >= 2:
    rows_data = [
        ('Contanti',  'Clicca sui pulsanti delle banconote e monete per comporre l\'importo. Il sistema calcola il resto automaticamente.'),
        ('Carta',     'Registra il pagamento con carta. L\'importo è automaticamente quello del totale.'),
        ('Satispay',  'Registra il pagamento tramite Satispay. Viene tracciato separatamente nelle statistiche. (Nuovo in v1.4)'),
        ('Omaggio',   'Registra l\'articolo come omaggio. Il totale diventa zero.'),
        ('Buono volontario', 'Usa il buono consumazione del volontario selezionato. Vedi sezione dedicata sopra. (Nuovo in v1.4)'),
    ]
    # Rimuovi righe esistenti tranne header
    while len(t_pag.rows) > 1:
        tr = t_pag.rows[-1]._tr
        t_pag._tbl.remove(tr)
    # Aggiungi righe nuove
    for metodo, descr in rows_data:
        row = t_pag.add_row()
        set_cell_text(row.cells[0], metodo)
        set_cell_text(row.cells[1], descr)

print("✓ Sezione 'Pagamento' aggiornata")

# ── 7. Rimuovi sezione "Sconti" ───────────────────────────────────────────────
p_sconti_h2 = find('Sconti', 'Heading20')
if p_sconti_h2:
    # Rimuovi H2 e i paragrafi seguenti fino al prossimo H2
    # Confronta per _element, non per identità wrapper
    all_paras = list(doc.paragraphs)
    sconti_el = p_sconti_h2._element
    idx = next((i for i, p in enumerate(all_paras) if p._element is sconti_el), None)
    to_delete = [p_sconti_h2]
    if idx is not None:
        for p in all_paras[idx+1:]:
            pPr = p._element.find(qn('w:pPr'))
            ps = pPr.find(qn('w:pStyle')) if pPr is not None else None
            sid = ps.get(qn('w:val')) if ps is not None else 'Normal'
            if sid in ('Heading10', 'Heading20'):
                break
        to_delete.append(p)
    for p in to_delete:
        try: delete_para(p)
        except: pass

print("✓ Sezione 'Sconti' rimossa")

# ── 8. Aggiorna "Storico e modifiche" ────────────────────────────────────────
p_storico_intro = find('Il pulsante Storico mostra gli scontrini della giornata')
if p_storico_intro:
    set_text(p_storico_intro,
        'Il pulsante Storico mostra gli scontrini della giornata. Da qui puoi:')

# Aggiunge "Modificare il tipo di pagamento" dopo "Stornare completamente"
p_storno = find('Stornare completamente uno scontrino', 'Paragrafoelenco')
if p_storno:
    insert_after(p_storno, tmpl_bullet,
        'Modificare il tipo di pagamento di uno scontrino già emesso (nuovo in v1.4)')

print("✓ 'Storico' aggiornato")

# ── 9. Aggiorna "8. Statistiche" ─────────────────────────────────────────────
p_stat_intro = find('La pagina Statistiche è organizzata in tre schede')
if p_stat_intro:
    set_text(p_stat_intro,
        'La pagina Statistiche è organizzata in tre schede: Corrente (scontrini della sessione aperta), Per data (filtrabile per intervallo), Sessioni archiviate (storico delle chiusure precedenti).')

p_filtro_h2 = find('Filtro per data (scheda Per data)', 'Heading20')
if p_filtro_h2:
    set_text(p_filtro_h2, 'Filtro per tipo di pagamento (nuovo in v1.4)')

p_filtro_testo = find('In cima alla pagina puoi impostare un intervallo di date')
if p_filtro_testo:
    set_text(p_filtro_testo,
        'In cima alla pagina sono disponibili pulsanti rapidi per filtrare gli scontrini per metodo di pagamento: Tutti · Contanti · Carta · Satispay · Buono · Omaggio.')

# Aggiorna tabella "I dati disponibili"
p_dati_h2 = find('I dati disponibili', 'Heading20')
t_dati = get_table_after(p_dati_h2) if p_dati_h2 else None
if t_dati:
    rows_dati = [
        ('Incasso reale',        'Soldi fisicamente incassati (contanti + carta + Satispay) — esclude buoni e omaggi'),
        ('Scontrini',            'Numero totale di scontrini validi'),
        ('Media scontrino',      'Valore medio per scontrino (calcolato sull\'incasso reale)'),
        ('Contanti',             'Totale pagato in contanti'),
        ('Carta',                'Totale pagato con carta'),
        ('Satispay',             'Totale pagato con Satispay (nuovo in v1.4)'),
        ('Buoni (costo oratorio)','Valore consumato dai volontari tramite buono — è un costo per l\'oratorio, non un\'entrata (nuovo in v1.4)'),
        ('Omaggi',               'Valore degli articoli omaggiati'),
        ('Pezzi venduti',        'Numero totale di articoli venduti'),
    ]
    while len(t_dati.rows) > 1:
        t_dati._tbl.remove(t_dati.rows[-1]._tr)
    for voce, descr in rows_dati:
        row = t_dati.add_row()
        set_cell_text(row.cells[0], voce)
        set_cell_text(row.cells[1], descr)

print("✓ Tabella 'I dati disponibili' aggiornata")

# Aggiorna Export Excel
p_export = find('Il pulsante Export XLSX genera un file Excel')
if p_export:
    set_text(p_export, 'Il pulsante ⬇ Excel genera un file Excel (.xlsx) con tre fogli:')

p_v_venduto  = find('Venduto: quantità e incasso per ogni prodotto')
if p_v_venduto: set_text(p_v_venduto,  'Venduto: quantità e incasso per ogni prodotto, inclusi gli omaggi')
p_v_scont    = find('Scontrini: elenco completo degli scontrini')
if p_v_scont:  set_text(p_v_scont,    'Scontrini: elenco completo con numero, data, lordo, buono usato, netto, tipo pagamento, stato')
p_v_riepi    = find('Riepilogo: totali aggregati per incasso')
if p_v_riepi:  set_text(p_v_riepi,   'Riepilogo: incasso reale, buoni (costo), scontrini, media, contanti, carta, Satispay, valore venduto totale')

print("✓ Export Excel aggiornato")

# Aggiorna Chiusura cassa
p_chius_intro = find('La chiusura cassa archivia la sessione corrente')
if p_chius_intro:
    set_text(p_chius_intro,
        'La chiusura cassa archivia la sessione corrente: tutti gli scontrini vengono collegati alla sessione e il contatore ricomincia da #0001. È una funzione riservata agli amministratori.')

p_ver_step = find("Verifica il riepilogo nell'anteprima: scontrini, totale", 'Paragrafoelenco')
if p_ver_step:
    set_text(p_ver_step,
        'Verifica il riepilogo nell\'anteprima: scontrini validi, incasso reale, buoni, contanti, carta, Satispay')

p_chius_gen  = find('Alla chiusura vengono generati automaticamente nella cartella chiusure')
if p_chius_gen:
    set_text(p_chius_gen,
        'Alla chiusura vengono generati automaticamente nella cartella chiusure\\:')

p_csv = find('File CSV compatibile Excel', 'Paragrafoelenco')
if p_csv:
    set_text(p_csv, 'File XLS multi-foglio (Riepilogo, Venduto, Scontrini) — aperto nativamente da Excel')

p_sug_email = find('Suggerimento: Se hai configurato i destinatari in Setup')
if p_sug_email:
    set_text(p_sug_email,
        'Suggerimento: Se hai configurato i destinatari in Setup → Notifiche, il report viene inviato via email con allegato XLS automaticamente alla chiusura.')

print("✓ 'Chiusura cassa' aggiornata")

# ── 10. Struttura dei file — aggiorna tabella ────────────────────────────────
p_struct_h1 = find('11. Struttura dei file', 'Heading10')
p_struct_norm = find('Questa e la struttura della cartella di Cassa Dalila')
t_struct = get_table_after(p_struct_norm) if p_struct_norm else None
if t_struct:
    # Aggiunge righe mancanti
    existing_texts = [r.cells[0].text.strip() for r in t_struct.rows]
    new_rows = [
        ('app/pb_migrations/', 'Script di inizializzazione e aggiornamento del database'),
        ('app/pb_hooks/',      'Logica server-side: scarico magazzino, chiusura cassa, invio email'),
        ('chiusure/',          'Report XLS e HTML generati ad ogni chiusura cassa'),
    ]
    for cell0_text, cell1_text in new_rows:
        if not any(cell0_text in t for t in existing_texts):
            row = t_struct.add_row()
            set_cell_text(row.cells[0], cell0_text)
            set_cell_text(row.cells[1], cell1_text)
    # Aggiorna riga backup\ se esiste
    for row in t_struct.rows:
        if 'backup' in row.cells[0].text and 'crea' not in row.cells[1].text:
            set_cell_text(row.cells[1], 'Cartella dei backup (creata automaticamente da BACKUP_DATI.bat)')

print("✓ Tabella 'Struttura dei file' aggiornata")

# ── 11. Aggiorna lista collection nel pannello admin ────────────────────────
p_collection_list = find('Ripeti per tutte le collection: scontrini, righe_scontrino', 'Paragrafoelenco')
if p_collection_list:
    set_text(p_collection_list,
        'Ripeti per tutte le collection: scontrini, righe_scontrino, prodotti, famiglie, comande, '
        'magazzini_comuni, movimenti_magazzino, comande_evase, utenti, menu, tavoli, configurazione, '
        'righe_pronte, volontari, sessioni_cassa')

print("✓ Lista collection aggiornata")

# ── Salva ─────────────────────────────────────────────────────────────────────
doc.save(DST)
print(f"\nDONE → {DST}")
